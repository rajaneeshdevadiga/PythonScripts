import docx

# Function to count occurrences of a specific word in a paragraph
def count_word_occurrences(paragraph, target_word):
    count = 0
    for word in paragraph.text.split():
        if word.lower() == target_word.lower():
            count += 1
    return count

# Function to search for specific words in a Word document and generate a report
def search_and_generate_report(docx_file, target_words):
    doc = docx.Document(docx_file)

    result = {}
    total_occurrences = 0

    for target_word in target_words:
        result[target_word] = 0

    for paragraph in doc.paragraphs:
        for target_word in target_words:
            count = count_word_occurrences(paragraph, target_word)
            result[target_word] += count
            total_occurrences += count

    # Generate the report
    report = f"Word Occurrence Report for '{docx_file}':\n"
    report += "-" * 50 + "\n"
    for target_word, count in result.items():
        report += f"'{target_word}': {count} occurrences\n"
    report += "-" * 50 + "\n"
    report += f"Total Occurrences: {total_occurrences}\n"

    return report

if __name__ == "__main__":
    # Prompt the user for the location of the Word document file
    docx_file_path = input("Enter the path to the Word document file: ")

    try:
        doc = docx.Document(docx_file_path)
    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        exit(1)

    # Prompt the user to enter words to search for
    target_words = []
    while True:
        search_word = input("Enter a word to search for (or press Enter to finish): ")
        if not search_word:
            break
        target_words.append(search_word)

    if not target_words:
        print("No words entered. Exiting...")
    else:
        report = search_and_generate_report(docx_file_path, target_words)

        # Print the report to the console
        print(report)

